from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from difflib import SequenceMatcher, get_close_matches
import fnmatch
import json
import os
from pathlib import Path
import re
import shutil
import sys
from typing import Dict, Iterable, List, Sequence, cast
from urllib.parse import urlparse


WIKI_DIRNAME = "wiki"
RAW_DIRNAME = "raw"
OUTPUTS_DIRNAME = "outputs"
CATEGORY_ORDER = ("core", "entities", "concepts", "analyses", "sources")
MARKDOWN_LINK_RE = re.compile(r"(?<!!)\[[^\]]+\]\(([^)]+)\)")
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".svg"}
EXTRACTABLE_SOURCE_SUFFIXES = {".pdf", ".docx", ".xlsx"}
MISSING_SUMMARY_TEXT = "Summary not yet written."
DEFAULT_PAGE_STATUS = "draft"
HOOK_HELPER_NAME = "llm-wiki-refresh"
HOOK_METADATA_NAME = "llm-wiki-hook.json"
HOOK_BACKUP_SUFFIX = ".llmwiki.bak"
HOOK_MARKER = "llm-wiki-managed"
MANAGED_HOOK_NAMES = ("pre-commit", "post-checkout", "post-merge")


@dataclass(frozen=True)
class LinkReference:
    source: Path
    target_text: str
    resolved: Path
    is_markdown_page: bool


@dataclass(frozen=True)
class Page:
    path: Path
    title: str
    category: str
    summary: str
    links: Sequence[LinkReference]


@dataclass(frozen=True)
class LintReport:
    broken_links: Sequence[LinkReference]
    orphan_pages: Sequence[Page]
    missing_summaries: Sequence[Page]

    @property
    def has_issues(self) -> bool:
        return bool(self.broken_links or self.orphan_pages or self.missing_summaries)


@dataclass(frozen=True)
class HealSuggestion:
    kind: str
    message: str


@dataclass(frozen=True)
class HealReport:
    suggestions: Sequence[HealSuggestion]

    @property
    def has_suggestions(self) -> bool:
        return bool(self.suggestions)


@dataclass(frozen=True)
class CompileItem:
    original_path: Path
    raw_source: Path
    wiki_page: Path
    kind: str


@dataclass(frozen=True)
class CompileReport:
    items: Sequence[CompileItem]

    @property
    def has_items(self) -> bool:
        return bool(self.items)


@dataclass(frozen=True)
class BriefReport:
    title: str
    content: str
    pages: Sequence[Page]
    output_path: Path | None


@dataclass(frozen=True)
class StatusReport:
    title: str
    content: str
    output_path: Path | None


@dataclass(frozen=True)
class WatchCycleReport:
    changed_paths: Sequence[Path]
    compiled_items: Sequence[CompileItem]
    index_rebuilt: bool
    status_output: Path | None
    messages: Sequence[str]

    @property
    def has_work(self) -> bool:
        return bool(self.compiled_items or self.index_rebuilt or self.status_output or self.messages)


@dataclass(frozen=True)
class GitHookInstallReport:
    git_dir: Path
    helper_path: Path
    metadata_path: Path
    hook_paths: Sequence[Path]
    backup_paths: Sequence[Path]
    status_output: str | None


@dataclass(frozen=True)
class GitHookStatus:
    git_dir: Path
    helper_path: Path
    metadata_path: Path
    helper_managed: bool
    metadata_managed: bool
    installed_hooks: Sequence[str]
    missing_hooks: Sequence[str]
    unmanaged_hooks: Sequence[str]
    backup_paths: Sequence[Path]
    status_output: str | None

    @property
    def is_installed(self) -> bool:
        return (
            self.helper_managed
            and self.metadata_managed
            and not self.missing_hooks
            and not self.unmanaged_hooks
            and len(self.installed_hooks) == len(MANAGED_HOOK_NAMES)
        )


@dataclass(frozen=True)
class GitHookUninstallReport:
    git_dir: Path
    removed_paths: Sequence[Path]
    restored_paths: Sequence[Path]


@dataclass(frozen=True)
class WebSource:
    url: str
    title: str
    markdown: str
    excerpt: str
    author: str | None
    published_date: str | None
    site_name: str | None
    fetched_at: str


def init_workspace(root: Path) -> List[Path]:
    root = root.resolve()
    created: List[Path] = []

    directories = [
        root / RAW_DIRNAME / "inbox",
        root / RAW_DIRNAME / "sources",
        root / RAW_DIRNAME / "assets",
        root / OUTPUTS_DIRNAME,
        root / WIKI_DIRNAME / "analyses",
        root / WIKI_DIRNAME / "concepts",
        root / WIKI_DIRNAME / "entities",
        root / WIKI_DIRNAME / "sources",
    ]

    for directory in directories:
        if not directory.exists():
            directory.mkdir(parents=True, exist_ok=True)
            created.append(directory)

    keep_files = [directory / ".gitkeep" for directory in directories]
    for keep_file in keep_files:
        if not keep_file.exists():
            keep_file.write_text("", encoding="utf-8")
            created.append(keep_file)

    seed_files = {
        root / WIKI_DIRNAME / "overview.md": (
            "# Overview\n\n"
            "This wiki is the durable synthesis layer between the raw sources and future questions.\n\n"
            "## Current Focus\n\n"
            "- Record the main themes, workstreams, or research questions guiding the wiki right now.\n\n"
            "## Priority Gaps\n\n"
            "- Track missing sources, unresolved questions, and contradictions worth closing next.\n\n"
            "## Related Pages\n\n"
            "- Link the concept, entity, and analysis pages that define the current map.\n"
        ),
        root / WIKI_DIRNAME / "log.md": (
            "# Log\n\n"
            "Append-only activity log. Each entry should start with a timestamped heading.\n"
        ),
        root / "AGENTS.md": render_agents_guidelines(),
        root / "CLAUDE.md": render_claude_schema(),
        root / "PROMPTS.md": render_prompt_library(),
        root / ".llmwikiignore": render_ignore_template(),
    }

    for path, content in seed_files.items():
        if not path.exists():
            path.write_text(content, encoding="utf-8")
            created.append(path)

    index_existed = (root / WIKI_DIRNAME / "index.md").exists()
    index_path = rebuild_index(root)
    if not index_existed:
        created.append(index_path)
    return created


def ingest_source(
    root: Path,
    source: Path,
    title: str | None = None,
    summary: str | None = None,
    kind: str = "auto",
    remove_source: bool = False,
    rebuild: bool = True,
) -> Dict[str, Path | str]:
    root = root.resolve()
    source = source.resolve()
    if not source.exists() or not source.is_file():
        raise FileNotFoundError(f"Source file not found: {source}")

    init_workspace(root)

    source_kind = classify_source_kind(source, kind)
    raw_bucket = "assets" if source_kind == "asset" else "sources"
    dated_name = datetime.now().strftime("%Y%m%d") + "-" + slugify(source.stem)
    raw_dest = unique_path(root / RAW_DIRNAME / raw_bucket, dated_name, source.suffix.lower())
    if remove_source:
        shutil.move(str(source), raw_dest)
    else:
        shutil.copy2(source, raw_dest)

    page_title = title or humanize_slug(source.stem)
    page_path = unique_path(root / WIKI_DIRNAME / "sources", slugify(page_title), ".md")
    raw_rel = relative_link(page_path, raw_dest)
    extracted_content = extract_structured_source_markdown(raw_dest) if source_kind == "source" else None
    extracted_summary = extract_summary(extracted_content) if extracted_content else ""
    page_summary = summary or extracted_summary or default_summary(page_title, source_kind)
    page_content = render_source_page(
        page_title,
        page_summary,
        raw_dest.name,
        raw_rel,
        source_kind,
        extracted_content=extracted_content,
    )
    page_path.write_text(page_content, encoding="utf-8")

    append_log_entry(
        root,
        kind="ingest",
        title=page_title,
        bullets=[
            f"Kind: `{source_kind}`",
            f"Raw source: [{raw_dest.name}]({relative_link(root / WIKI_DIRNAME / 'log.md', raw_dest)})",
            f"Wiki page: [{page_title}]({relative_link(root / WIKI_DIRNAME / 'log.md', page_path)})",
            "Status: source page created; summarize it and connect it to the rest of the wiki.",
        ],
    )
    if rebuild:
        rebuild_index(root)
    return {"raw_source": raw_dest, "wiki_page": page_path, "kind": source_kind}


def ingest_url(
    root: Path,
    url: str,
    title: str | None = None,
    summary: str | None = None,
    rebuild: bool = True,
) -> Dict[str, Path | str]:
    root = root.resolve()
    init_workspace(root)
    validate_web_url(url)

    web_source = fetch_web_source(url)
    page_title = title or web_source.title or fallback_title_from_url(url)
    raw_dest = unique_path(
        root / RAW_DIRNAME / "sources",
        datetime.now().strftime("%Y%m%d") + "-" + slugify(page_title),
        ".md",
    )
    raw_dest.write_text(render_web_source_snapshot(web_source, page_title), encoding="utf-8")

    page_path = unique_path(root / WIKI_DIRNAME / "sources", slugify(page_title), ".md")
    page_summary = summary or web_source.excerpt or default_summary(page_title, "source")
    raw_rel = relative_link(page_path, raw_dest)
    source_lines = [
        f"- Raw file: [{raw_dest.name}]({raw_rel})",
        "- Kind: `source`",
        f"- URL: [{web_source.url}]({web_source.url})",
        f"- Site: `{web_source.site_name}`" if web_source.site_name else "",
        f"- Author: `{web_source.author}`" if web_source.author else "",
        f"- Published: `{web_source.published_date}`" if web_source.published_date else "",
        f"- Fetched: `{web_source.fetched_at}`",
    ]
    page_content = render_source_page(
        page_title,
        page_summary,
        raw_dest.name,
        raw_rel,
        "source",
        source_lines=[line for line in source_lines if line],
    )
    page_path.write_text(page_content, encoding="utf-8")

    append_log_entry(
        root,
        kind="ingest",
        title=page_title,
        bullets=[
            "Kind: `source`",
            f"URL: [{web_source.url}]({web_source.url})",
            f"Raw source: [{raw_dest.name}]({relative_link(root / WIKI_DIRNAME / 'log.md', raw_dest)})",
            f"Wiki page: [{page_title}]({relative_link(root / WIKI_DIRNAME / 'log.md', page_path)})",
            "Status: web source ingested; review the extraction and connect it to the rest of the wiki.",
        ],
    )
    if rebuild:
        rebuild_index(root)
    return {"raw_source": raw_dest, "wiki_page": page_path, "kind": "source", "url": web_source.url}


def create_page(root: Path, title: str, category: str, summary: str | None = None) -> Path:
    root = root.resolve()
    init_workspace(root)

    category_dir = root / WIKI_DIRNAME / category
    category_dir.mkdir(parents=True, exist_ok=True)
    keep_file = category_dir / ".gitkeep"
    if not keep_file.exists():
        keep_file.write_text("", encoding="utf-8")

    page_path = unique_path(category_dir, slugify(title), ".md")
    page_summary = summary or f"Starter page for {title}. Write a concise synthesis before treating it as settled."
    content = (
        f"{render_frontmatter(title, source_count=0)}\n\n"
        f"# {title}\n\n"
        f"{page_summary}\n\n"
        "## Key Points\n\n"
        "- Add the main ideas or findings.\n\n"
        "## Evidence\n\n"
        "- Add supporting facts, examples, or citations in `[Source: page-or-file]` form.\n\n"
        "## Contradictions\n\n"
        "- Note disagreements, stale claims, or superseded interpretations.\n\n"
        "## Related Pages\n\n"
        "- Link the most relevant wiki pages.\n\n"
        "## Open Questions\n\n"
        "- Record unresolved questions or ambiguities.\n"
    )
    page_path.write_text(content, encoding="utf-8")

    append_log_entry(
        root,
        kind="page",
        title=title,
        bullets=[
            f"Category: `{category}`",
            f"Wiki page: [{title}]({relative_link(root / WIKI_DIRNAME / 'log.md', page_path)})",
        ],
    )
    rebuild_index(root)
    return page_path


def rebuild_index(root: Path) -> Path:
    root = root.resolve()
    wiki_dir = root / WIKI_DIRNAME
    wiki_dir.mkdir(parents=True, exist_ok=True)
    index_path = wiki_dir / "index.md"

    pages = [page for page in scan_pages(root) if page.path.name not in {"index.md", "log.md"}]
    grouped: Dict[str, List[Page]] = {}
    for page in pages:
        grouped.setdefault(page.category, []).append(page)

    lines = [
        "# Index",
        "",
        "Read this page first when answering questions. It is the content-oriented map of the wiki.",
        "",
        "Regenerate with `python3 -m llm_wiki index` after structural edits.",
    ]

    ordered_categories = list(CATEGORY_ORDER) + sorted(
        category for category in grouped if category not in CATEGORY_ORDER
    )
    for category in ordered_categories:
        category_pages = sorted(grouped.get(category, []), key=lambda item: item.title.lower())
        if not category_pages:
            continue
        lines.extend(["", f"## {category.title()}"])
        for page in category_pages:
            target = relative_link(index_path, page.path)
            summary = clip(page.summary or MISSING_SUMMARY_TEXT, 160)
            lines.append(f"- [{page.title}]({target}) - {summary}")

    index_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return index_path


def scan_pages(root: Path) -> List[Page]:
    root = root.resolve()
    wiki_dir = root / WIKI_DIRNAME
    if not wiki_dir.exists():
        return []

    pages: List[Page] = []
    for path in sorted(wiki_dir.rglob("*.md")):
        text = path.read_text(encoding="utf-8")
        title = extract_title(text) or humanize_slug(path.stem)
        summary = extract_summary(text)
        category = classify_page(path, wiki_dir)
        links = parse_links(path, text)
        pages.append(Page(path=path.resolve(), title=title, category=category, summary=summary, links=links))
    return pages


def lint_workspace(root: Path) -> LintReport:
    root = root.resolve()
    pages = scan_pages(root)
    wiki_dir = root / WIKI_DIRNAME
    index_path = (wiki_dir / "index.md").resolve()
    log_path = (wiki_dir / "log.md").resolve()
    overview_path = (wiki_dir / "overview.md").resolve()

    broken_links: List[LinkReference] = []
    inbound: Dict[Path, set[Path]] = {}
    missing_summaries: List[Page] = []

    for page in pages:
        if page.path.name not in {"index.md", "log.md"} and not page.summary:
            missing_summaries.append(page)

        for link in page.links:
            if is_within(link.resolved, root) and not link.resolved.exists():
                broken_links.append(link)

            if not link.is_markdown_page:
                continue
            if not is_within(link.resolved, wiki_dir):
                continue
            if not link.resolved.exists():
                continue
            if page.path in {index_path, log_path}:
                continue
            inbound.setdefault(link.resolved, set()).add(page.path)

    orphan_pages = [
        page
        for page in pages
        if page.path.name not in {"index.md", "log.md"}
        and page.path != overview_path
        and not inbound.get(page.path)
    ]

    return LintReport(
        broken_links=sorted(broken_links, key=lambda item: (str(item.source), item.target_text)),
        orphan_pages=sorted(orphan_pages, key=lambda item: str(item.path)),
        missing_summaries=sorted(missing_summaries, key=lambda item: str(item.path)),
    )


def search_pages(root: Path, query: str, limit: int = 10) -> List[tuple[int, Page]]:
    query_terms = tokenize(query)
    if not query_terms:
        return []

    results: List[tuple[int, Page]] = []
    for page in scan_pages(root):
        if page.path.name in {"index.md", "log.md"}:
            continue
        text = page.path.read_text(encoding="utf-8").lower()
        title_text = page.title.lower()
        summary_text = page.summary.lower()
        score = 0
        for term in query_terms:
            score += title_text.count(term) * 8
            score += summary_text.count(term) * 4
            score += text.count(term)
        if " ".join(query_terms) in title_text:
            score += 10
        if score > 0:
            results.append((score, page))

    results.sort(key=lambda item: (-item[0], item[1].title.lower()))
    return results[:limit]


def compile_inbox(root: Path, limit: int | None = None) -> CompileReport:
    root = root.resolve()
    init_workspace(root)

    inbox_dir = root / RAW_DIRNAME / "inbox"
    inbox_files = list_inbox_files(inbox_dir)
    if limit is not None:
        inbox_files = inbox_files[: max(limit, 0)]

    items: List[CompileItem] = []
    for source in inbox_files:
        result = ingest_source(root, source, remove_source=True, rebuild=False)
        raw_source = cast(Path, result["raw_source"])
        wiki_page = cast(Path, result["wiki_page"])
        items.append(
            CompileItem(
                original_path=source,
                raw_source=raw_source,
                wiki_page=wiki_page,
                kind=str(result["kind"]),
            )
        )

    prune_empty_dirs(inbox_dir)
    if items:
        append_log_entry(
            root,
            kind="compile",
            title="inbox batch",
            bullets=[
                f"Processed {len(items)} inbox files into canonical raw storage and wiki source pages.",
                "Review the new source pages, summarize them, and connect them to the rest of the wiki.",
            ],
        )
        rebuild_index(root)
    return CompileReport(items=items)


def process_watch_changes(
    root: Path,
    changed_paths: Sequence[str | Path],
    *,
    status_output_path: Path | None = None,
    limit: int = 5,
) -> WatchCycleReport:
    root = root.resolve()
    relevant = filter_watch_paths(root, changed_paths, status_output_path=status_output_path)
    if not relevant:
        return WatchCycleReport(
            changed_paths=[],
            compiled_items=[],
            index_rebuilt=False,
            status_output=None,
            messages=[],
        )

    inbox_changed = any(is_within(path, root / RAW_DIRNAME / "inbox") for path in relevant)
    wiki_changed = any(is_within(path, root / WIKI_DIRNAME) for path in relevant)
    compiled_items: List[CompileItem] = []
    index_rebuilt = False
    status_output: Path | None = None
    messages: List[str] = []

    if inbox_changed:
        compile_report = compile_inbox(root)
        compiled_items = list(compile_report.items)
        if compiled_items:
            messages.append(f"Compiled {len(compiled_items)} inbox file(s).")
            index_rebuilt = True
        else:
            messages.append("Inbox changed, but there were no compilable files.")

    if wiki_changed and not compiled_items:
        rebuild_index(root)
        index_rebuilt = True
        messages.append("Rebuilt wiki index after markdown changes.")

    if status_output_path is not None and (inbox_changed or wiki_changed):
        status_report = build_status_report(root, output_path=status_output_path, limit=limit)
        status_output = status_report.output_path
        if status_output is not None:
            messages.append(f"Wrote status report to {status_output}.")

    return WatchCycleReport(
        changed_paths=relevant,
        compiled_items=compiled_items,
        index_rebuilt=index_rebuilt,
        status_output=status_output,
        messages=messages,
    )


def watch_workspace(
    root: Path,
    *,
    status_output_path: Path | None = None,
    limit: int = 5,
    debounce_ms: int = 1000,
) -> None:
    try:
        from watchfiles import watch
    except ImportError as exc:
        raise RuntimeError(
            "Watch mode requires `watchfiles`. Install it with `python3 -m pip install 'llm-wiki[watch]'`."
        ) from exc

    root = root.resolve()
    init_workspace(root)

    initial_inbox = list_inbox_files(root / RAW_DIRNAME / "inbox")
    if initial_inbox:
        startup_report = process_watch_changes(
            root,
            initial_inbox,
            status_output_path=status_output_path,
            limit=limit,
        )
        emit_watch_messages(startup_report)

    for changes in watch(root, debounce=debounce_ms, raise_interrupt=True):
        changed_paths = [Path(path_text).resolve() for _, path_text in changes]
        report = process_watch_changes(
            root,
            changed_paths,
            status_output_path=status_output_path,
            limit=limit,
        )
        emit_watch_messages(report)


def install_git_hooks(
    root: Path,
    *,
    status_output_path: Path | None = None,
    force: bool = False,
) -> GitHookInstallReport:
    root = root.resolve()
    init_workspace(root)
    git_dir = resolve_git_dir(root)
    hooks_dir = git_dir / "hooks"
    hooks_dir.mkdir(parents=True, exist_ok=True)

    configured_status_output = normalize_hook_status_output(root, status_output_path)
    staged_status_output = hook_stage_path(root, configured_status_output)
    helper_path = hooks_dir / HOOK_HELPER_NAME
    metadata_path = hooks_dir / HOOK_METADATA_NAME
    hook_paths = [hooks_dir / name for name in MANAGED_HOOK_NAMES]
    targets = [helper_path, metadata_path, *hook_paths]
    backup_paths: List[Path] = []
    conflicts: List[Path] = []

    for target in targets:
        if not target.exists():
            continue
        if is_managed_hook_path(target):
            continue
        if not force:
            conflicts.append(target)
            continue
        backup_paths.append(backup_hook_path(target))

    if conflicts:
        formatted = ", ".join(str(path) for path in conflicts)
        raise RuntimeError(
            f"Refusing to overwrite unmanaged git hook files: {formatted}. "
            "Re-run with `--force` to back them up and replace them."
        )

    metadata = {
        "managed_by": "llm-wiki",
        "root": str(root),
        "python": sys.executable,
        "pythonpath": str(module_search_root()),
        "status_output": configured_status_output,
        "stage_status_output": staged_status_output,
        "hooks": list(MANAGED_HOOK_NAMES),
    }

    write_executable_script(
        helper_path,
        render_hook_helper(
            root=root,
            python_executable=sys.executable,
            pythonpath=module_search_root(),
            status_output=configured_status_output,
            staged_status_output=staged_status_output,
        ),
    )
    metadata_path.write_text(json.dumps(metadata, indent=2) + "\n", encoding="utf-8")
    for hook_name, hook_path in zip(MANAGED_HOOK_NAMES, hook_paths):
        write_executable_script(hook_path, render_hook_wrapper(hook_name))

    return GitHookInstallReport(
        git_dir=git_dir,
        helper_path=helper_path,
        metadata_path=metadata_path,
        hook_paths=hook_paths,
        backup_paths=backup_paths,
        status_output=configured_status_output,
    )


def inspect_git_hooks(root: Path) -> GitHookStatus:
    root = root.resolve()
    git_dir = resolve_git_dir(root)
    hooks_dir = git_dir / "hooks"
    helper_path = hooks_dir / HOOK_HELPER_NAME
    metadata_path = hooks_dir / HOOK_METADATA_NAME
    metadata = read_hook_metadata(metadata_path)

    installed_hooks: List[str] = []
    missing_hooks: List[str] = []
    unmanaged_hooks: List[str] = []
    for hook_name in MANAGED_HOOK_NAMES:
        hook_path = hooks_dir / hook_name
        if not hook_path.exists():
            missing_hooks.append(hook_name)
            continue
        if is_managed_hook_path(hook_path):
            installed_hooks.append(hook_name)
        else:
            unmanaged_hooks.append(hook_name)

    return GitHookStatus(
        git_dir=git_dir,
        helper_path=helper_path,
        metadata_path=metadata_path,
        helper_managed=is_managed_hook_path(helper_path),
        metadata_managed=metadata is not None,
        installed_hooks=tuple(installed_hooks),
        missing_hooks=tuple(missing_hooks),
        unmanaged_hooks=tuple(unmanaged_hooks),
        backup_paths=tuple(find_hook_backups(hooks_dir)),
        status_output=str(metadata.get("status_output")) if metadata and metadata.get("status_output") else None,
    )


def uninstall_git_hooks(root: Path) -> GitHookUninstallReport:
    root = root.resolve()
    git_dir = resolve_git_dir(root)
    hooks_dir = git_dir / "hooks"
    helper_path = hooks_dir / HOOK_HELPER_NAME
    metadata_path = hooks_dir / HOOK_METADATA_NAME

    removed_paths: List[Path] = []
    restored_paths: List[Path] = []
    for path in [helper_path, metadata_path, *(hooks_dir / name for name in MANAGED_HOOK_NAMES)]:
        if path.exists() and is_managed_hook_path(path):
            path.unlink()
            removed_paths.append(path)

    for path in [helper_path, metadata_path, *(hooks_dir / name for name in MANAGED_HOOK_NAMES)]:
        restored = restore_hook_backup(path)
        if restored is not None:
            restored_paths.append(restored)

    return GitHookUninstallReport(
        git_dir=git_dir,
        removed_paths=removed_paths,
        restored_paths=restored_paths,
    )


def build_brief(
    root: Path,
    query: str | None = None,
    page_refs: Sequence[str] = (),
    limit: int = 5,
    title: str | None = None,
    output_path: Path | None = None,
) -> BriefReport:
    root = root.resolve()
    pages = resolve_brief_pages(root, query=query, page_refs=page_refs, limit=limit)
    if not pages:
        raise ValueError("No wiki pages matched the brief request.")

    resolved_output = resolve_output_path(root, output_path) if output_path else None
    report_title = title or default_brief_title(query, pages)
    content = render_brief(root, report_title, pages, query=query, output_path=resolved_output)

    if resolved_output:
        resolved_output.parent.mkdir(parents=True, exist_ok=True)
        resolved_output.write_text(content, encoding="utf-8")
        if is_within(resolved_output, root / WIKI_DIRNAME):
            append_log_entry(
                root,
                kind="brief",
                title=report_title,
                bullets=[
                    f"Report: [{report_title}]({relative_link(root / WIKI_DIRNAME / 'log.md', resolved_output)})",
                    f"Pages used: {len(pages)}",
                    f"Query: `{query}`" if query else "Query: explicit page selection",
                ],
            )
            rebuild_index(root)

    return BriefReport(title=report_title, content=content, pages=pages, output_path=resolved_output)


def build_status_report(
    root: Path,
    title: str | None = None,
    output_path: Path | None = None,
    limit: int = 5,
) -> StatusReport:
    root = root.resolve()
    pages = scan_pages(root)
    content_pages = [page for page in pages if page.path.name not in {"index.md", "log.md"}]
    lint = lint_workspace(root)
    heal = heal_workspace(root)
    duplicates = find_duplicate_candidates(content_pages)
    raw_files = collect_raw_files(root)
    raw_counts = count_raw_files_by_bucket(root, raw_files)
    output_files = list_output_files(root)
    source_pages = [page for page in content_pages if page.category == "sources"]
    synthesis_pages = [page for page in content_pages if page.category in {"analyses", "concepts", "entities"}]
    citation_gaps = find_pages_missing_citations(synthesis_pages)
    category_counts = Counter(page.category for page in content_pages)
    status_counts = Counter(extract_page_status(page.path.read_text(encoding="utf-8")) for page in content_pages)
    recent_entries = read_recent_log_entries(root, limit=limit)
    unindexed_raw = find_unindexed_raw_files(root, pages)
    linked_raw_count = max(len(raw_files) - len(unindexed_raw), 0)
    cited_synthesis_count = max(len(synthesis_pages) - len(citation_gaps), 0)
    raw_word_count = sum(estimate_word_count(path) for path in raw_files)

    report_title = title or f"{humanize_slug(root.name)} Status"
    resolved_output = resolve_output_path(root, output_path) if output_path else None
    content = render_status_report(
        root=root,
        title=report_title,
        raw_counts=raw_counts,
        raw_word_count=raw_word_count,
        content_pages=content_pages,
        category_counts=category_counts,
        status_counts=status_counts,
        output_files=output_files,
        source_pages=source_pages,
        linked_raw_count=linked_raw_count,
        synthesis_pages=synthesis_pages,
        cited_synthesis_count=cited_synthesis_count,
        lint=lint,
        duplicates=duplicates,
        citation_gaps=citation_gaps,
        unindexed_raw=unindexed_raw,
        heal=heal,
        recent_entries=recent_entries,
        limit=limit,
    )

    if resolved_output:
        resolved_output.parent.mkdir(parents=True, exist_ok=True)
        resolved_output.write_text(content, encoding="utf-8")

    return StatusReport(title=report_title, content=content, output_path=resolved_output)


def heal_workspace(root: Path) -> HealReport:
    root = root.resolve()
    pages = scan_pages(root)
    lint = lint_workspace(root)
    suggestions: List[HealSuggestion] = []
    wiki_dir = root / WIKI_DIRNAME

    for item in lint.broken_links:
        path_guess = suggest_closest_page(item.target_text, pages)
        if path_guess:
            rel_guess = relative_link(item.source, path_guess.path)
            message = (
                f"{item.source.name} links to missing target `{item.target_text}`. "
                f"Closest existing page: [{path_guess.title}]({rel_guess})."
            )
        else:
            message = (
                f"{item.source.name} links to missing target `{item.target_text}`. "
                "Either create the page or fix the link."
            )
        suggestions.append(HealSuggestion(kind="broken_link", message=message))

    for page in lint.orphan_pages:
        related = suggest_related_pages(page, pages, limit=2)
        if related:
            options = ", ".join(
                f"[{candidate.title}]({relative_link(page.path, candidate.path)})" for candidate in related
            )
            message = f"{page.title} is orphaned. Link it from {options}, or from `overview.md`."
        else:
            message = f"{page.title} is orphaned. Add a link from `overview.md` or another core page."
        suggestions.append(HealSuggestion(kind="orphan_page", message=message))

    for page in lint.missing_summaries:
        message = (
            f"{relative_link(wiki_dir / 'index.md', page.path)} is missing an opening summary paragraph. "
            "Add a concise first paragraph so search and indexing stay useful."
        )
        suggestions.append(HealSuggestion(kind="missing_summary", message=message))

    for left, right, score in find_duplicate_candidates(pages):
        message = (
            f"Possible duplicate topics: {left.title} and {right.title} "
            f"(similarity {score:.2f}). Consider merging or cross-linking them."
        )
        suggestions.append(HealSuggestion(kind="duplicate_candidate", message=message))

    for raw_file in find_unindexed_raw_files(root, pages):
        target_dir = "raw/inbox" if is_within(raw_file, root / RAW_DIRNAME / "inbox") else f"raw/{raw_file.parent.name}"
        message = (
            f"Raw file `{raw_file.name}` in `{target_dir}` has no matching wiki source page. "
            f"Ingest it with `python3 -m llm_wiki ingest {raw_file}`."
        )
        suggestions.append(HealSuggestion(kind="unindexed_raw", message=message))

    suggestions.sort(key=lambda item: (heal_priority(item.kind), item.message.lower()))
    return HealReport(suggestions=suggestions)


def append_log_entry(root: Path, kind: str, title: str, bullets: Iterable[str]) -> Path:
    log_path = root / WIKI_DIRNAME / "log.md"
    log_path.parent.mkdir(parents=True, exist_ok=True)
    if not log_path.exists():
        log_path.write_text("# Log\n\n", encoding="utf-8")

    stamp = datetime.now().astimezone().strftime("%Y-%m-%d %H:%M")
    lines = [f"## [{stamp}] {kind} | {title}", ""]
    for bullet in bullets:
        lines.append(f"- {bullet}")
    lines.append("")

    existing = log_path.read_text(encoding="utf-8")
    suffix = "\n" if existing and not existing.endswith("\n\n") else ""
    log_path.write_text(existing + suffix + "\n".join(lines), encoding="utf-8")
    return log_path


def parse_links(page_path: Path, text: str) -> List[LinkReference]:
    links: List[LinkReference] = []
    for target in MARKDOWN_LINK_RE.findall(text):
        cleaned = target.strip().split("#", 1)[0]
        if not cleaned or re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", cleaned):
            continue
        resolved = (page_path.parent / cleaned).resolve()
        links.append(
            LinkReference(
                source=page_path.resolve(),
                target_text=target.strip(),
                resolved=resolved,
                is_markdown_page=cleaned.endswith(".md"),
            )
        )
    return links


def render_source_page(
    title: str,
    summary: str,
    raw_name: str,
    raw_rel: str,
    source_kind: str,
    source_lines: Sequence[str] = (),
    extracted_content: str | None = None,
) -> str:
    return render_source_page_with_metadata(
        title=title,
        summary=summary,
        raw_name=raw_name,
        raw_rel=raw_rel,
        source_kind=source_kind,
        source_lines=source_lines,
        extracted_content=extracted_content,
    )


def render_source_page_with_metadata(
    title: str,
    summary: str,
    raw_name: str,
    raw_rel: str,
    source_kind: str,
    source_lines: Sequence[str],
    extracted_content: str | None = None,
) -> str:
    lines = [
        render_frontmatter(title, source_count=1),
        "",
        f"# {title}",
        "",
        summary,
        "",
        "## Source",
        "",
    ]
    if source_lines:
        lines.extend(source_lines)
    else:
        lines.extend(
            [
                f"- Raw file: [{raw_name}]({raw_rel})",
                f"- Kind: `{source_kind}`",
            ]
        )
    lines.append("")

    if source_kind == "asset":
        lines.extend(
            [
                "## Preview",
                "",
                f"![{title}]({raw_rel})",
                "",
                "## Visual Notes",
                "",
                "- Describe the important visual structures, labels, and relationships.",
                "",
            ]
        )
    elif extracted_content:
        lines.extend(
            [
                "## Extracted Content",
                "",
                extracted_content.strip(),
                "",
            ]
        )

    lines.extend(
        [
            "## Key Points",
            "",
            "- Capture the main takeaways from this source.",
            "",
            "## Evidence",
            "",
            f"- When carrying claims into synthesis pages, cite this source as `[Source: {raw_name}]`.",
            "",
            "## Contradictions",
            "",
            "- Record where this source conflicts with the current wiki.",
            "",
            "## Related Pages",
            "",
            "- Link the wiki pages this source informs.",
            "",
            "## Open Questions",
            "",
            "- Record what still needs clarification.",
            "",
        ]
    )
    return "\n".join(lines)


def render_brief(
    root: Path,
    title: str,
    pages: Sequence[Page],
    query: str | None = None,
    output_path: Path | None = None,
) -> str:
    timestamp = datetime.now().astimezone().strftime("%Y-%m-%d %H:%M %Z")
    intro = (
        f"This brief compiles {len(pages)} wiki pages relevant to `{query}`."
        if query
        else f"This brief compiles {len(pages)} selected wiki pages."
    )

    lines = [f"# {title}", "", intro, "", "## Scope", "", f"- Generated: {timestamp}"]
    if query:
        lines.append(f"- Query: `{query}`")
    lines.append(f"- Pages: {len(pages)}")

    lines.extend(["", "## Highlights", ""])
    for page in pages:
        lines.append(f"- {page.title}: {page.summary or MISSING_SUMMARY_TEXT}")

    lines.extend(["", "## Source Pages", ""])
    for page in pages:
        lines.append(f"- {format_page_reference(page, root, output_path)} - {page.summary or MISSING_SUMMARY_TEXT}")

    lines.extend(["", "## Notes By Page"])
    for page in pages:
        lines.extend(["", f"### {page.title}", "", page.summary or MISSING_SUMMARY_TEXT, ""])
        lines.append(f"Source: {format_page_reference(page, root, output_path)}")
        points = extract_notable_points(page.path.read_text(encoding="utf-8"), limit=4)
        if points:
            lines.extend(["", "Key details:", ""])
            for point in points:
                lines.append(f"- {point}")

    lines.extend(["", "## Related Pages", ""])
    for page in pages:
        lines.append(f"- {format_page_reference(page, root, output_path)}")

    return "\n".join(lines).rstrip() + "\n"


def render_status_report(
    *,
    root: Path,
    title: str,
    raw_counts: Counter[str],
    raw_word_count: int,
    content_pages: Sequence[Page],
    category_counts: Counter[str],
    status_counts: Counter[str],
    output_files: Sequence[Path],
    source_pages: Sequence[Page],
    linked_raw_count: int,
    synthesis_pages: Sequence[Page],
    cited_synthesis_count: int,
    lint: LintReport,
    duplicates: Sequence[tuple[Page, Page, float]],
    citation_gaps: Sequence[Page],
    unindexed_raw: Sequence[Path],
    heal: HealReport,
    recent_entries: Sequence[str],
    limit: int,
) -> str:
    timestamp = datetime.now().astimezone().strftime("%Y-%m-%d %H:%M %Z")
    total_raw = sum(raw_counts.values())
    verdict = workspace_verdict(
        total_raw=total_raw,
        content_pages=len(content_pages),
        lint=lint,
        unindexed_raw=len(unindexed_raw),
    )

    lines = [
        f"# {title}",
        "",
        f"A workspace health snapshot for `{root.name}` generated {timestamp}.",
        "",
        "## Corpus Check",
        "",
        f"- Raw files: {total_raw} ({format_bucket_counts(raw_counts)})",
        f"- Approximate raw text words: {raw_word_count:,}",
        f"- Wiki pages: {len(content_pages)} content pages ({format_counter(category_counts, CATEGORY_ORDER)})",
        f"- Saved outputs: {len(output_files)}",
        f"- Verdict: {verdict}",
        "",
        "## Coverage",
        "",
        f"- Raw files linked from the wiki: {linked_raw_count}/{total_raw}" if total_raw else "- Raw files linked from the wiki: 0/0",
        f"- Source pages: {len(source_pages)}",
        (
            f"- Synthesis pages with `[Source: ...]` citations: {cited_synthesis_count}/{len(synthesis_pages)}"
            if synthesis_pages
            else "- Synthesis pages with `[Source: ...]` citations: 0/0"
        ),
        f"- Unindexed raw files: {len(unindexed_raw)}",
        "",
        "## Structural Issues",
        "",
        f"- Broken links: {len(lint.broken_links)}",
        f"- Orphan pages: {len(lint.orphan_pages)}",
        f"- Missing summaries: {len(lint.missing_summaries)}",
        f"- Duplicate topic candidates: {len(duplicates)}",
        f"- Citation gaps: {len(citation_gaps)}",
        f"- Heal suggestions: {len(heal.suggestions)}",
        "",
        "## Page Status",
        "",
        f"- {format_counter(status_counts, ('draft', 'reviewed', 'needs_update', 'unknown'))}",
    ]

    lines.extend(["", "## Recent Activity", ""])
    if recent_entries:
        for entry in recent_entries:
            lines.append(f"- {entry}")
    else:
        lines.append("- No log entries yet.")

    lines.extend(["", "## Knowledge Gaps", ""])
    append_gap_items(lines, "Unindexed raw files", [str(path.relative_to(root)) for path in unindexed_raw], limit)
    append_gap_items(lines, "Pages missing citations", [page.title for page in citation_gaps], limit)
    append_gap_items(lines, "Orphan pages", [page.title for page in lint.orphan_pages], limit)
    append_gap_items(
        lines,
        "Duplicate candidates",
        [f"{left.title} <> {right.title} ({score:.2f})" for left, right, score in duplicates],
        limit,
    )
    if lines[-1] == "":
        lines.append("- No immediate gap signals found.")

    lines.extend(["", "## Suggested Actions", ""])
    if heal.suggestions:
        for suggestion in heal.suggestions[: max(limit, 0)]:
            lines.append(f"- [{suggestion.kind}] {suggestion.message}")
    else:
        lines.append("- No immediate repair actions. Next leverage: ingest more material or turn recurring questions into analysis pages.")

    return "\n".join(lines).rstrip() + "\n"


def extract_title(text: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped[2:].strip()
    return ""


def extract_summary(text: str) -> str:
    text = strip_frontmatter(text)
    paragraph: List[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            if paragraph:
                break
            continue
        if stripped.startswith("#") or stripped.startswith(">") or stripped.startswith("```"):
            if paragraph:
                break
            continue
        if stripped.startswith("- ") or stripped.startswith("* ") or re.match(r"^\d+\.\s", stripped):
            if paragraph:
                break
            continue
        paragraph.append(stripped)
    return " ".join(paragraph).strip()


def strip_frontmatter(text: str) -> str:
    if not text.startswith("---\n"):
        return text
    parts = text.split("\n---\n", 1)
    if len(parts) == 2:
        return parts[1]
    return text


def classify_page(path: Path, wiki_dir: Path) -> str:
    relative = path.resolve().relative_to(wiki_dir.resolve())
    if len(relative.parts) == 1:
        return "core"
    return relative.parts[0]


def relative_link(from_path: Path, to_path: Path) -> str:
    rel = os.path.relpath(to_path, start=from_path.parent)
    return rel.replace(os.sep, "/")


def unique_path(directory: Path, stem: str, suffix: str) -> Path:
    directory.mkdir(parents=True, exist_ok=True)
    candidate = directory / f"{stem}{suffix}"
    counter = 2
    while candidate.exists():
        candidate = directory / f"{stem}-{counter}{suffix}"
        counter += 1
    return candidate


def slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip().lower()).strip("-")
    return slug or "untitled"


def humanize_slug(value: str) -> str:
    value = value.replace("-", " ").replace("_", " ").strip()
    return re.sub(r"\s+", " ", value).title() or "Untitled"


def tokenize(value: str) -> List[str]:
    return re.findall(r"[a-z0-9]+", value.lower())


def clip(value: str, limit: int) -> str:
    if len(value) <= limit:
        return value
    return value[: limit - 3].rstrip() + "..."


def is_within(path: Path, root: Path) -> bool:
    try:
        path.resolve().relative_to(root.resolve())
        return True
    except ValueError:
        return False


def classify_source_kind(source: Path, requested_kind: str) -> str:
    if requested_kind not in {"auto", "source", "asset"}:
        raise ValueError(f"Unsupported kind: {requested_kind}")
    if requested_kind != "auto":
        return requested_kind
    return "asset" if source.suffix.lower() in IMAGE_SUFFIXES else "source"


def default_summary(title: str, source_kind: str) -> str:
    if source_kind == "asset":
        return f"Starter visual source page for {title}. Describe what the image shows and why it matters."
    return f"Starter source page for {title}. Summarize the file after reviewing it."


def extract_structured_source_markdown(path: Path) -> str | None:
    suffix = path.suffix.lower()
    if suffix not in EXTRACTABLE_SOURCE_SUFFIXES:
        return None
    if suffix == ".pdf":
        return extract_pdf_markdown(path)
    if suffix == ".docx":
        return extract_docx_markdown(path)
    if suffix == ".xlsx":
        return extract_xlsx_markdown(path)
    return None


def extract_pdf_markdown(path: Path) -> str | None:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "PDF ingest requires `pypdf`. Install it with `python3 -m pip install 'llm-wiki[office]'`."
        ) from exc

    reader = PdfReader(str(path))
    pages: List[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        normalized = normalize_extracted_text(text)
        if not normalized:
            continue
        pages.append(f"## Page {index}\n\n{normalized}")
    return "\n\n".join(pages).strip() or None


def extract_docx_markdown(path: Path) -> str | None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX ingest requires `python-docx`. Install it with `python3 -m pip install 'llm-wiki[office]'`."
        ) from exc

    doc = Document(str(path))
    lines: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        elif style.startswith("List"):
            lines.append(f"- {text}")
        else:
            lines.append(text)
    for table in doc.tables:
        lines.extend(render_docx_table(table))
    content = "\n".join(lines).strip()
    return content or None


def render_docx_table(table: object) -> List[str]:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return []
    lines = [
        "",
        "| " + " | ".join(rows[0]) + " |",
        "| " + " | ".join("---" for _ in rows[0]) + " |",
    ]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    return lines


def extract_xlsx_markdown(path: Path) -> str | None:
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError(
            "XLSX ingest requires `openpyxl`. Install it with `python3 -m pip install 'llm-wiki[office]'`."
        ) from exc

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    try:
        sections: List[str] = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            rows = [
                ["" if cell is None else str(cell) for cell in row]
                for row in sheet.iter_rows(values_only=True)
                if any(cell is not None for cell in row)
            ]
            if not rows:
                continue
            sections.append(f"## Sheet: {sheet_name}")
            sections.append("")
            sections.append("| " + " | ".join(rows[0]) + " |")
            sections.append("| " + " | ".join("---" for _ in rows[0]) + " |")
            for row in rows[1:]:
                sections.append("| " + " | ".join(row) + " |")
            sections.append("")
        content = "\n".join(sections).strip()
        return content or None
    finally:
        workbook.close()


def normalize_extracted_text(text: str) -> str:
    cleaned_lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    non_empty = [line for line in cleaned_lines if line]
    return "\n\n".join(non_empty)


def validate_web_url(url: str) -> None:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise ValueError(f"Unsupported URL: {url}")


def fetch_web_source(url: str) -> WebSource:
    try:
        import trafilatura
    except ImportError as exc:
        raise RuntimeError(
            "Web ingest requires `trafilatura`. Install it with `python3 -m pip install trafilatura`."
        ) from exc

    downloaded = trafilatura.fetch_url(url)
    if not downloaded:
        raise ValueError(f"Unable to fetch URL: {url}")

    metadata = trafilatura.bare_extraction(downloaded, url=url, with_metadata=True, as_dict=True) or {}
    markdown = trafilatura.extract(
        downloaded,
        url=url,
        output_format="markdown",
        with_metadata=False,
        include_links=True,
        include_images=False,
        include_comments=False,
    )
    body = (markdown or metadata.get("text") or "").strip()
    if not body:
        raise ValueError(f"Unable to extract readable content from URL: {url}")

    title = str(metadata.get("title") or fallback_title_from_url(url))
    excerpt_source = str(metadata.get("description") or body)
    excerpt = clip(extract_summary(excerpt_source) or excerpt_source.strip(), 280)
    fetched_at = datetime.now().astimezone().strftime("%Y-%m-%d %H:%M %Z")
    return WebSource(
        url=str(metadata.get("url") or url),
        title=title,
        markdown=body,
        excerpt=excerpt,
        author=clean_optional_string(metadata.get("author")),
        published_date=clean_optional_string(metadata.get("date")),
        site_name=clean_optional_string(metadata.get("sitename") or metadata.get("hostname")),
        fetched_at=fetched_at,
    )


def render_web_source_snapshot(source: WebSource, title: str) -> str:
    metadata_lines = [
        "---",
        f'title: "{yaml_quote(title)}"',
        f'url: "{yaml_quote(source.url)}"',
        f'fetched_at: "{yaml_quote(source.fetched_at)}"',
    ]
    if source.site_name:
        metadata_lines.append(f'site_name: "{yaml_quote(source.site_name)}"')
    if source.author:
        metadata_lines.append(f'author: "{yaml_quote(source.author)}"')
    if source.published_date:
        metadata_lines.append(f'published_date: "{yaml_quote(source.published_date)}"')
    metadata_lines.extend(["---", "", f"# {title}", "", f"Source URL: {source.url}", ""])

    if source.author or source.published_date or source.site_name:
        metadata_lines.extend(["## Metadata", ""])
        if source.site_name:
            metadata_lines.append(f"- Site: `{source.site_name}`")
        if source.author:
            metadata_lines.append(f"- Author: `{source.author}`")
        if source.published_date:
            metadata_lines.append(f"- Published: `{source.published_date}`")
        metadata_lines.extend(["", "## Extracted Content", ""])

    metadata_lines.extend([source.markdown.strip(), ""])
    return "\n".join(metadata_lines)


def fallback_title_from_url(url: str) -> str:
    parsed = urlparse(url)
    tail = Path(parsed.path.rstrip("/")).name
    candidate = tail or parsed.netloc or url
    return humanize_slug(candidate)


def clean_optional_string(value: object) -> str | None:
    if not isinstance(value, str):
        return None
    cleaned = value.strip()
    return cleaned or None


def render_frontmatter(title: str, source_count: int, status: str = DEFAULT_PAGE_STATUS) -> str:
    stamp = datetime.now().astimezone().strftime("%Y-%m-%d")
    return "\n".join(
        [
            "---",
            f'title: "{yaml_quote(title)}"',
            f"created: {stamp}",
            f"last_updated: {stamp}",
            f"source_count: {source_count}",
            f"status: {status}",
            "---",
        ]
    )


def yaml_quote(value: str) -> str:
    return value.replace("\\", "\\\\").replace('"', '\\"')


def render_agents_guidelines() -> str:
    return (
        "# LLM Wiki Agent Guidelines\n\n"
        "This repo is a persistent thinking layer. Treat it like a disciplined markdown knowledge base, not like a chat transcript.\n\n"
        "## Operating Model\n\n"
        "- `raw/` stores immutable source material.\n"
        "- `wiki/` stores the compiled markdown knowledge base.\n"
        "- `outputs/` stores generated briefs, reports, and durable answer artifacts.\n"
        "- `wiki/index.md` is the content map.\n"
        "- `wiki/log.md` is the chronological change trail.\n\n"
        "## Directory Rules\n\n"
        "- Never edit files in `raw/` after they have been ingested or compiled.\n"
        "- Put images, diagrams, screenshots, and other visual artifacts in `raw/assets/`.\n"
        "- Keep source interpretation in `wiki/sources/`.\n"
        "- Keep reusable ideas in `wiki/concepts/`.\n"
        "- Keep people, organizations, projects, datasets, and tools in `wiki/entities/`.\n"
        "- Keep comparisons, deep dives, recurring questions, and durable reports in `wiki/analyses/`.\n"
        "- Use relative markdown links for internal references.\n\n"
        "## Core Loop\n\n"
        "1. Ingest new material with `python3 -m llm_wiki ingest ...`, `python3 -m llm_wiki ingest-url ...`, or batch-process `raw/inbox/` with `python3 -m llm_wiki compile`.\n"
        "2. Replace generated starter text with a sharp first paragraph, key points, evidence, contradictions, and open questions.\n"
        "3. Pull durable ideas into concept, entity, and analysis pages instead of leaving them stranded in source pages.\n"
        "4. Read `wiki/index.md` and use `python3 -m llm_wiki search \"<query>\"` before answering.\n"
        "5. File durable answers back into the wiki, and write shareable outputs to `outputs/` when useful.\n"
        "6. Run `python3 -m llm_wiki watch --status-output outputs/status.md` while editing, or install managed git hooks with `python3 -m llm_wiki hook install --status-output outputs/status.md` so commit, checkout, and merge flows keep derived files current.\n\n"
        "## Commands\n\n"
        "- `hook install`: install managed git hooks for pre-commit, post-checkout, and post-merge refreshes.\n"
        "- `hook status`: show whether the managed hooks are installed cleanly.\n"
        "- `hook uninstall`: remove managed hooks and restore any backups created with `--force`.\n\n"
        "## Page Rules\n\n"
        "- Every page gets exactly one `# Title`.\n"
        "- Generated pages include YAML frontmatter for title, dates, source count, and status.\n"
        "- The first paragraph must stand alone as the page summary. It drives indexing and retrieval.\n"
        "- Cite synthesized claims in `[Source: page-or-file]` form whenever possible.\n"
        "- Keep a `## Contradictions` section current when sources disagree or older claims have been superseded.\n"
        "- Keep a `## Related Pages` section current whenever a page touches existing material.\n"
        "- If a question keeps recurring, convert it into a durable analysis page.\n\n"
        "## Style\n\n"
        "- Keep markdown terse and scannable.\n"
        "- Prefer bullets for evidence, distinctions, and open questions.\n"
        "- Preserve uncertainty explicitly when sources disagree or remain incomplete.\n"
        "- Do not hide contradictions; surface them and link the conflicting pages.\n"
    )


def render_claude_schema() -> str:
    return (
        "# Knowledge Base Schema\n\n"
        "## Identity\n\n"
        "This is a persistent markdown knowledge base maintained by an LLM agent. The human curates sources, directs the analysis, and asks questions. The LLM handles ingest, synthesis, linking, maintenance, and report generation.\n\n"
        "## Architecture\n\n"
        "- `raw/` contains immutable source documents. Never modify files in `raw/` after ingest.\n"
        "- `wiki/` contains the compiled wiki.\n"
        "- `outputs/` contains generated reports, analyses, and saved answers.\n"
        "- `AGENTS.md` and `PROMPTS.md` define the operating rules and reusable prompts.\n"
        "- Managed git hooks can refresh `wiki/index.md` and a status snapshot during commit, checkout, and merge flows.\n\n"
        "## Wiki Conventions\n\n"
        "- Use one durable page per topic.\n"
        "- Generated pages start with YAML frontmatter for `title`, `created`, `last_updated`, `source_count`, and `status`.\n"
        "- Keep a one-paragraph summary immediately under the `# Title` heading.\n"
        "- Use normal markdown links for internal references.\n"
        "- Cite synthesized claims in `[Source: page-or-file]` form.\n"
        "- Track disagreements under `## Contradictions`.\n"
        "- Keep `## Related Pages` current.\n\n"
        "## Workflows\n\n"
        "- Ingest: read local files or fetched web snapshots, summarize them, update or create relevant wiki pages, add backlinks, note contradictions, and log the change.\n"
        "- Status: generate a concise health snapshot before major synthesis passes so the agent sees coverage gaps, structural issues, and next actions.\n"
        "- Query: read `wiki/index.md` first, then read relevant pages before answering. Save durable outputs when they are worth keeping.\n"
        "- Lint: check for broken links, orphan pages, duplicate topics, missing summaries, stale claims, and missing citations.\n"
        "- Brief: package existing wiki knowledge into a clean markdown artifact in `outputs/`.\n"
        "- Hook automation: if managed hooks are installed, let them refresh derived files instead of maintaining `wiki/index.md` or `outputs/status.md` by hand.\n"
    )


def render_prompt_library() -> str:
    return (
        "# Prompt Library\n\n"
        "Use these prompts with a local-file-aware coding agent. If your agent prefers `CLAUDE.md`, reference that file instead of `AGENTS.md`.\n\n"
        "## Ingest One Source\n\n"
        "```text\n"
        "Read AGENTS.md and CLAUDE.md. Process [FILENAME] from raw/ or raw/inbox/. Read it fully, then create or update the relevant wiki pages, add backlinks, note contradictions, update wiki/index.md, and append to wiki/log.md.\n"
        "```\n\n"
        "## Ingest URL\n\n"
        "```text\n"
        "Read AGENTS.md and CLAUDE.md. Fetch [URL], review the extracted markdown snapshot, then create or update the relevant wiki pages, add backlinks, note contradictions, update wiki/index.md, and append to wiki/log.md.\n"
        "```\n\n"
        "## Compile Inbox\n\n"
        "```text\n"
        "Read AGENTS.md and CLAUDE.md. Process all unprocessed files in raw/inbox/ sequentially. For each file, create or update the relevant source page, connect it to the rest of the wiki, update wiki/index.md, and log the work.\n"
        "```\n\n"
        "## Query\n\n"
        "```text\n"
        "Read wiki/index.md first. Based on the existing wiki, answer: [QUESTION]. Cite the wiki pages or source pages that informed the answer. If the answer reveals a durable new synthesis, file it back into wiki/ or outputs/.\n"
        "```\n\n"
        "## Explore\n\n"
        "```text\n"
        "Read wiki/index.md and identify the 5 most interesting unexplored connections between existing topics. For each, explain the possible insight, the evidence already in the wiki, and what additional source would help confirm it.\n"
        "```\n\n"
        "## Brief\n\n"
        "```text\n"
        "Read wiki/index.md and the most relevant pages. Write a concise markdown brief on [TOPIC] using only wiki knowledge, cite the pages used, and save the result in outputs/[topic]-brief.md.\n"
        "```\n\n"
        "## Status Snapshot\n\n"
        "```text\n"
        "Run `python3 -m llm_wiki status --output outputs/status.md` and review the biggest structural risks, coverage gaps, and next actions before making major changes.\n"
        "```\n\n"
        "## Watch Mode\n\n"
        "```text\n"
        "Start `python3 -m llm_wiki watch --status-output outputs/status.md` while editing the wiki so inbox drops compile automatically, the index stays fresh, and the current workspace status is always available.\n"
        "```\n\n"
        "## Hook Setup\n\n"
        "```text\n"
        "Install managed git hooks with `python3 -m llm_wiki hook install --status-output outputs/status.md` so pre-commit, post-checkout, and post-merge keep derived files current.\n"
        "```\n\n"
        "## Lint\n\n"
        "```text\n"
        "Run a health check on wiki/. Identify broken links, orphan pages, duplicate topics, missing summaries, stale claims, contradictions, and unsupported claims without clear source attribution. Save the findings in outputs/lint-report-[date].md.\n"
        "```\n"
    )


def suggest_closest_page(target_text: str, pages: Sequence[Page]) -> Page | None:
    target_name = Path(target_text.split("#", 1)[0]).stem
    candidates = {
        page.title.lower(): page for page in pages if page.path.name not in {"index.md", "log.md"}
    }
    candidates.update(
        {page.path.stem.lower(): page for page in pages if page.path.name not in {"index.md", "log.md"}}
    )
    matches = get_close_matches(target_name.lower(), list(candidates.keys()), n=1, cutoff=0.6)
    if not matches:
        return None
    return candidates[matches[0]]


def suggest_related_pages(page: Page, pages: Sequence[Page], limit: int = 2) -> List[Page]:
    ranked: List[tuple[int, Page]] = []
    own_terms = set(tokenize(page.title + " " + page.summary))
    for candidate in pages:
        if candidate.path == page.path or candidate.path.name in {"index.md", "log.md"}:
            continue
        candidate_terms = set(tokenize(candidate.title + " " + candidate.summary))
        overlap = len(own_terms & candidate_terms)
        score = overlap * 10
        if candidate.category == page.category:
            score += 2
        if score > 0:
            ranked.append((score, candidate))
    ranked.sort(key=lambda item: (-item[0], item[1].title.lower()))
    return [candidate for _, candidate in ranked[:limit]]


def find_duplicate_candidates(pages: Sequence[Page]) -> List[tuple[Page, Page, float]]:
    candidates: List[tuple[Page, Page, float]] = []
    content_pages = [page for page in pages if page.path.name not in {"index.md", "log.md"}]
    for index, left in enumerate(content_pages):
        left_key = f"{left.title} {left.summary}".lower()
        for right in content_pages[index + 1 :]:
            right_key = f"{right.title} {right.summary}".lower()
            score = SequenceMatcher(None, left_key, right_key).ratio()
            if score >= 0.74:
                candidates.append((left, right, score))
    candidates.sort(key=lambda item: (-item[2], item[0].title.lower(), item[1].title.lower()))
    return candidates[:5]


def find_unindexed_raw_files(root: Path, pages: Sequence[Page]) -> List[Path]:
    raw_dir = root / RAW_DIRNAME
    if not raw_dir.exists():
        return []
    ignore_patterns = load_ignore_patterns(root)

    referenced: set[Path] = set()
    for page in pages:
        for link in page.links:
            if is_within(link.resolved, raw_dir) and link.resolved.exists():
                referenced.add(link.resolved)

    raw_files: List[Path] = []
    for directory in (raw_dir / "sources", raw_dir / "assets", raw_dir / "inbox"):
        if not directory.exists():
            continue
        for path in sorted(directory.rglob("*")):
            if not path.is_file() or path.name == ".gitkeep":
                continue
            if is_ignored(path.resolve(), root, ignore_patterns):
                continue
            resolved = path.resolve()
            if resolved not in referenced:
                raw_files.append(resolved)
    return raw_files


def heal_priority(kind: str) -> int:
    order = {
        "broken_link": 0,
        "missing_summary": 1,
        "orphan_page": 2,
        "duplicate_candidate": 3,
        "unindexed_raw": 4,
    }
    return order.get(kind, 99)


def list_inbox_files(inbox_dir: Path) -> List[Path]:
    if not inbox_dir.exists():
        return []
    root = inbox_dir.parent.parent.resolve()
    ignore_patterns = load_ignore_patterns(root)
    return [
        path.resolve()
        for path in sorted(inbox_dir.rglob("*"))
        if path.is_file() and path.name != ".gitkeep"
        if not is_ignored(path.resolve(), root, ignore_patterns)
    ]


def prune_empty_dirs(directory: Path) -> None:
    if not directory.exists():
        return
    for path in sorted((item for item in directory.rglob("*") if item.is_dir()), reverse=True):
        if any(path.iterdir()):
            continue
        path.rmdir()


def resolve_brief_pages(root: Path, query: str | None, page_refs: Sequence[str], limit: int) -> List[Page]:
    if not query and not page_refs:
        raise ValueError("Provide a query or at least one page reference for the brief.")

    pages = [page for page in scan_pages(root) if page.path.name not in {"index.md", "log.md"}]
    selected: List[Page] = []
    seen: set[Path] = set()

    for ref in page_refs:
        page = resolve_page_reference(ref, pages, root)
        if page is None:
            raise ValueError(f"No wiki page matched '{ref}'.")
        if page.path not in seen:
            selected.append(page)
            seen.add(page.path)

    if query:
        for _, page in search_pages(root, query, limit=limit):
            if page.path not in seen:
                selected.append(page)
                seen.add(page.path)

    return selected


def resolve_page_reference(ref: str, pages: Sequence[Page], root: Path) -> Page | None:
    page_map = {page.path.resolve(): page for page in pages}
    ref_path = Path(ref)
    candidates: List[Path] = []
    if ref_path.is_absolute():
        candidates.append(ref_path.resolve())
    else:
        candidates.append((root / ref_path).resolve())
        candidates.append((root / WIKI_DIRNAME / ref_path).resolve())

    for candidate in candidates:
        if candidate in page_map:
            return page_map[candidate]

    ref_lower = ref.strip().lower()
    for page in pages:
        root_rel = str(page.path.relative_to(root)).replace(os.sep, "/").lower()
        wiki_rel = str(page.path.relative_to(root / WIKI_DIRNAME)).replace(os.sep, "/").lower()
        if ref_lower in {page.title.lower(), page.path.stem.lower(), root_rel, wiki_rel}:
            return page
    return None


def resolve_output_path(root: Path, output_path: Path) -> Path:
    if output_path.is_absolute():
        return output_path.resolve()
    return (root / output_path).resolve()


def default_brief_title(query: str | None, pages: Sequence[Page]) -> str:
    if query:
        return f"{humanize_slug(slugify(query))} Brief"
    if len(pages) == 1:
        return f"{pages[0].title} Brief"
    return "Wiki Brief"


def format_page_reference(page: Page, root: Path, output_path: Path | None) -> str:
    if output_path:
        return f"[{page.title}]({relative_link(output_path, page.path)})"
    rel = str(page.path.relative_to(root)).replace(os.sep, "/")
    return f"{page.title} (`{rel}`)"


def extract_notable_points(text: str, limit: int = 4) -> List[str]:
    preferred_sections = {
        "key points",
        "evidence",
        "visual notes",
        "architecture at a glance",
        "operating loop",
        "design implications for this repo",
    }
    ignored_sections = {"source", "related pages", "open questions"}
    preferred: List[str] = []
    fallback: List[str] = []
    current_section = ""

    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            current_section = stripped[3:].strip().lower()
            continue

        item = parse_list_item(stripped)
        if not item or current_section in ignored_sections:
            continue

        if current_section in preferred_sections:
            preferred.append(item)
        else:
            fallback.append(item)

    ordered = preferred + fallback
    unique: List[str] = []
    seen: set[str] = set()
    for item in ordered:
        if item in seen:
            continue
        unique.append(item)
        seen.add(item)
        if len(unique) >= limit:
            break
    return unique


def parse_list_item(line: str) -> str:
    if line.startswith("- ") or line.startswith("* "):
        return line[2:].strip()
    numbered = re.match(r"^\d+\.\s+(.*)$", line)
    if numbered:
        return numbered.group(1).strip()
    return ""


def load_ignore_patterns(root: Path) -> List[str]:
    ignore_file = root / ".llmwikiignore"
    if not ignore_file.exists():
        return []
    patterns: List[str] = []
    for line in ignore_file.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped and not stripped.startswith("#"):
            patterns.append(stripped)
    return patterns


def is_ignored(path: Path, root: Path, patterns: Sequence[str]) -> bool:
    if not patterns:
        return False
    try:
        rel = str(path.resolve().relative_to(root.resolve())).replace(os.sep, "/")
    except ValueError:
        return False

    parts = rel.split("/")
    for pattern in patterns:
        normalized = pattern.strip()
        if not normalized:
            continue
        trimmed = normalized.strip("/")
        if not trimmed:
            continue
        if normalized.endswith("/") and (rel == trimmed or rel.startswith(f"{trimmed}/")):
            return True
        if fnmatch.fnmatch(rel, trimmed):
            return True
        if fnmatch.fnmatch(path.name, trimmed):
            return True
        for index, part in enumerate(parts):
            if fnmatch.fnmatch(part, trimmed):
                return True
            if fnmatch.fnmatch("/".join(parts[: index + 1]), trimmed):
                return True
    return False


def render_ignore_template() -> str:
    return (
        "# Paths to skip during `compile` and raw-file health checks.\n"
        "# Syntax is path-oriented glob matching relative to the workspace root.\n"
        "# Examples:\n"
        "# raw/inbox/archive/\n"
        "# raw/inbox/*.tmp\n"
        "# raw/assets/screenshots/\n"
    )


def resolve_git_dir(root: Path) -> Path:
    dot_git = root / ".git"
    if dot_git.is_dir():
        return dot_git.resolve()
    if dot_git.is_file():
        content = dot_git.read_text(encoding="utf-8").strip()
        if content.startswith("gitdir:"):
            return (root / content.split(":", 1)[1].strip()).resolve()
    raise RuntimeError(
        f"No git repository found at {root}. Run `git init` before installing managed hooks."
    )


def normalize_hook_status_output(root: Path, status_output_path: Path | None) -> str | None:
    if status_output_path is None:
        return None
    resolved = resolve_output_path(root, status_output_path).resolve()
    if is_within(resolved, root):
        return str(resolved.relative_to(root)).replace(os.sep, "/")
    return str(resolved)


def hook_stage_path(root: Path, configured_status_output: str | None) -> str | None:
    if not configured_status_output:
        return None
    candidate = Path(configured_status_output)
    resolved = candidate.resolve() if candidate.is_absolute() else (root / candidate).resolve()
    if not is_within(resolved, root):
        return None
    return str(resolved.relative_to(root)).replace(os.sep, "/")


def render_hook_wrapper(hook_name: str) -> str:
    return "\n".join(
        [
            "#!/bin/sh",
            f"# {HOOK_MARKER}",
            f"LLM_WIKI_HOOK_NAME={shell_quote(hook_name)}",
            "export LLM_WIKI_HOOK_NAME",
            f'. "$(dirname "$0")/{HOOK_HELPER_NAME}"',
            "",
        ]
    )


def render_hook_helper(
    *,
    root: Path,
    python_executable: str,
    pythonpath: Path,
    status_output: str | None,
    staged_status_output: str | None,
) -> str:
    return "\n".join(
        [
            "#!/bin/sh",
            f"# {HOOK_MARKER}",
            "set -eu",
            f"LLM_WIKI_ROOT={shell_quote(str(root))}",
            f"LLM_WIKI_PYTHON={shell_quote(python_executable)}",
            f"LLM_WIKI_PYTHONPATH={shell_quote(str(pythonpath))}",
            f"LLM_WIKI_STATUS_OUTPUT={shell_quote(status_output or '')}",
            f"LLM_WIKI_STAGE_STATUS={shell_quote(staged_status_output or '')}",
            "",
            'LLM_WIKI_HOOK_NAME="${1:-${LLM_WIKI_HOOK_NAME:-}}"',
            'cd "$LLM_WIKI_ROOT"',
            'export PYTHONPATH="$LLM_WIKI_PYTHONPATH${PYTHONPATH:+:$PYTHONPATH}"',
            '"$LLM_WIKI_PYTHON" -m llm_wiki index --root "$LLM_WIKI_ROOT" >/dev/null',
            'if [ -n "$LLM_WIKI_STATUS_OUTPUT" ]; then',
            '  "$LLM_WIKI_PYTHON" -m llm_wiki status --root "$LLM_WIKI_ROOT" --output "$LLM_WIKI_STATUS_OUTPUT" >/dev/null',
            "fi",
            'if [ "$LLM_WIKI_HOOK_NAME" = "pre-commit" ]; then',
            '  git add -- "wiki/index.md"',
            '  if [ -n "$LLM_WIKI_STAGE_STATUS" ]; then',
            '    if git ls-files --error-unmatch -- "$LLM_WIKI_STAGE_STATUS" >/dev/null 2>&1 || ! git check-ignore -q -- "$LLM_WIKI_STAGE_STATUS"; then',
            '      git add -- "$LLM_WIKI_STAGE_STATUS"',
            "    fi",
            "  fi",
            "fi",
            "",
        ]
    )


def shell_quote(value: str) -> str:
    return "'" + value.replace("'", "'\"'\"'") + "'"


def module_search_root() -> Path:
    return Path(__file__).resolve().parent.parent


def write_executable_script(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")
    path.chmod(0o755)


def is_managed_hook_path(path: Path) -> bool:
    if not path.exists() or not path.is_file():
        return False
    if path.name == HOOK_METADATA_NAME:
        return read_hook_metadata(path) is not None
    try:
        text = path.read_text(encoding="utf-8")
    except OSError:
        return False
    return HOOK_MARKER in text


def read_hook_metadata(path: Path) -> Dict[str, object] | None:
    if not path.exists() or not path.is_file():
        return None
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    if not isinstance(payload, dict):
        return None
    if payload.get("managed_by") != "llm-wiki":
        return None
    return payload


def backup_hook_path(path: Path) -> Path:
    backup = unique_path(path.parent, f"{path.name}{HOOK_BACKUP_SUFFIX}", "")
    shutil.move(str(path), backup)
    return backup


def restore_hook_backup(path: Path) -> Path | None:
    backups = sorted(path.parent.glob(f"{path.name}{HOOK_BACKUP_SUFFIX}*"))
    if not backups or path.exists():
        return None
    backup = backups[-1]
    shutil.move(str(backup), path)
    return path


def find_hook_backups(hooks_dir: Path) -> List[Path]:
    if not hooks_dir.exists():
        return []
    return sorted(path.resolve() for path in hooks_dir.glob(f"*{HOOK_BACKUP_SUFFIX}*") if path.is_file())


def collect_raw_files(root: Path) -> List[Path]:
    raw_dir = root / RAW_DIRNAME
    if not raw_dir.exists():
        return []
    ignore_patterns = load_ignore_patterns(root)
    files: List[Path] = []
    for directory in (raw_dir / "sources", raw_dir / "assets", raw_dir / "inbox"):
        if not directory.exists():
            continue
        for path in sorted(directory.rglob("*")):
            if not path.is_file() or path.name == ".gitkeep":
                continue
            resolved = path.resolve()
            if is_ignored(resolved, root, ignore_patterns):
                continue
            files.append(resolved)
    return files


def count_raw_files_by_bucket(root: Path, raw_files: Sequence[Path]) -> Counter[str]:
    counts: Counter[str] = Counter()
    raw_root = (root / RAW_DIRNAME).resolve()
    for path in raw_files:
        try:
            relative = path.resolve().relative_to(raw_root)
        except ValueError:
            continue
        bucket = relative.parts[0] if relative.parts else "unknown"
        counts[bucket] += 1
    return counts


def list_output_files(root: Path) -> List[Path]:
    outputs_dir = root / OUTPUTS_DIRNAME
    if not outputs_dir.exists():
        return []
    return [
        path.resolve()
        for path in sorted(outputs_dir.rglob("*"))
        if path.is_file() and path.name != ".gitkeep"
    ]


def extract_frontmatter_fields(text: str) -> Dict[str, str]:
    if not text.startswith("---\n"):
        return {}
    parts = text.split("\n---\n", 1)
    if len(parts) != 2:
        return {}
    fields: Dict[str, str] = {}
    for line in parts[0].splitlines()[1:]:
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        fields[key.strip()] = value.strip().strip('"').strip("'")
    return fields


def extract_page_status(text: str) -> str:
    return extract_frontmatter_fields(text).get("status", "unknown") or "unknown"


def find_pages_missing_citations(pages: Sequence[Page]) -> List[Page]:
    missing: List[Page] = []
    for page in pages:
        text = page.path.read_text(encoding="utf-8")
        citations = re.findall(r"\[Source:\s*([^\]]+)\]", text)
        real_citations = [citation.strip() for citation in citations if citation.strip().lower() != "page-or-file"]
        if not real_citations:
            missing.append(page)
    return sorted(missing, key=lambda item: item.title.lower())


def read_recent_log_entries(root: Path, limit: int = 5) -> List[str]:
    log_path = root / WIKI_DIRNAME / "log.md"
    if not log_path.exists():
        return []
    entries: List[str] = []
    for line in log_path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            entries.append(stripped[3:])
    return entries[-max(limit, 0):][::-1]


def estimate_word_count(path: Path) -> int:
    if path.suffix.lower() in IMAGE_SUFFIXES:
        return 0
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
    except OSError:
        return 0
    return len(re.findall(r"[A-Za-z0-9_]+", text))


def workspace_verdict(total_raw: int, content_pages: int, lint: LintReport, unindexed_raw: int) -> str:
    if total_raw == 0 and content_pages == 0:
        return "Workspace is initialized but empty. Add sources before expecting useful synthesis."
    if total_raw < 5 and content_pages < 10:
        return "Thin corpus. The workflow works, but compounding value will stay low until you add more source material."
    if lint.has_issues or unindexed_raw:
        return "Usable but incomplete. Structural cleanup or pending raw ingest will improve answer quality."
    return "Coherent working set. The wiki is in good shape for search, briefs, and agent-guided synthesis."


def format_bucket_counts(counts: Counter[str]) -> str:
    order = ("sources", "assets", "inbox")
    parts = [f"{bucket} {counts.get(bucket, 0)}" for bucket in order if counts.get(bucket, 0)]
    if not parts:
        return "sources 0, assets 0, inbox 0"
    return ", ".join(parts)


def format_counter(counts: Counter[str], order: Sequence[str]) -> str:
    seen: set[str] = set()
    parts: List[str] = []
    for key in order:
        if key in counts:
            parts.append(f"{key} {counts[key]}")
            seen.add(key)
    for key in sorted(counts):
        if key in seen:
            continue
        parts.append(f"{key} {counts[key]}")
    return ", ".join(parts) if parts else "none"


def append_gap_items(lines: List[str], label: str, items: Sequence[str], limit: int) -> None:
    if not items:
        return
    preview = ", ".join(items[: max(limit, 0)])
    suffix = f" (+{len(items) - limit} more)" if len(items) > limit and limit > 0 else ""
    lines.append(f"- {label}: {preview}{suffix}")


def filter_watch_paths(
    root: Path,
    changed_paths: Sequence[str | Path],
    *,
    status_output_path: Path | None = None,
) -> List[Path]:
    root = root.resolve()
    resolved_status = resolve_output_path(root, status_output_path).resolve() if status_output_path else None
    relevant: List[Path] = []
    seen: set[Path] = set()

    for changed in changed_paths:
        path = Path(changed).resolve()
        if path in seen:
            continue
        if path.name == ".gitkeep":
            continue
        if resolved_status is not None and path == resolved_status:
            continue
        if is_within(path, root / OUTPUTS_DIRNAME):
            continue
        if is_within(path, root / WIKI_DIRNAME):
            if path.name in {"index.md", "log.md"}:
                continue
            if path.suffix.lower() != ".md":
                continue
            relevant.append(path)
            seen.add(path)
            continue
        if is_within(path, root / RAW_DIRNAME / "inbox"):
            relevant.append(path)
            seen.add(path)
            continue
    return relevant


def emit_watch_messages(report: WatchCycleReport) -> None:
    if not report.has_work:
        return
    stamp = datetime.now().astimezone().strftime("%Y-%m-%d %H:%M:%S %Z")
    print(f"[{stamp}] watch")
    for path in report.changed_paths[:10]:
        print(f"- changed: {path}")
    if len(report.changed_paths) > 10:
        print(f"- changed: ... and {len(report.changed_paths) - 10} more path(s)")
    for message in report.messages:
        print(f"- {message}")
