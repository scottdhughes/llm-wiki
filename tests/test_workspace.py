import tempfile
from pathlib import Path
import subprocess
import unittest
from unittest.mock import patch

from llm_wiki.workspace import (
    build_brief,
    build_status_report,
    compile_inbox,
    create_page,
    heal_workspace,
    ingest_url,
    ingest_source,
    inspect_git_hooks,
    install_git_hooks,
    init_workspace,
    lint_workspace,
    process_watch_changes,
    rebuild_index,
    uninstall_git_hooks,
    WebSource,
)


class WorkspaceTests(unittest.TestCase):
    def test_init_workspace_creates_outputs_and_compatibility_files(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)

            init_workspace(root)

            self.assertTrue((root / "outputs").is_dir())
            self.assertTrue((root / "outputs" / ".gitkeep").exists())
            self.assertTrue((root / "AGENTS.md").exists())
            self.assertTrue((root / "CLAUDE.md").exists())
            self.assertTrue((root / "PROMPTS.md").exists())
            self.assertTrue((root / ".llmwikiignore").exists())
            self.assertIn("outputs/", (root / "AGENTS.md").read_text(encoding="utf-8"))
            self.assertIn("## Ingest One Source", (root / "PROMPTS.md").read_text(encoding="utf-8"))

    def test_rebuild_index_lists_pages(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)
            create_page(root, "Attention", "concepts", "Mechanism for weighting relevant context.")
            create_page(root, "Competitor Matrix", "analyses", "Comparison of the main alternatives.")
            rebuild_index(root)

            index = (root / "wiki" / "index.md").read_text(encoding="utf-8")
            self.assertIn("[Attention](concepts/attention.md)", index)
            self.assertIn("Mechanism for weighting relevant context.", index)
            self.assertIn("[Competitor Matrix](analyses/competitor-matrix.md)", index)

    def test_ingest_copies_source_and_creates_source_page(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "notes.txt"
            source.write_text("hello world", encoding="utf-8")

            result = ingest_source(root, source, title="Sample Notes")

            raw_source = result["raw_source"]
            wiki_page = result["wiki_page"]
            self.assertTrue(raw_source.exists())
            self.assertTrue(wiki_page.exists())
            page_text = wiki_page.read_text(encoding="utf-8")
            self.assertIn('title: "Sample Notes"', page_text)
            self.assertIn("source_count: 1", page_text)
            self.assertIn("Sample Notes", page_text)
            self.assertIn(raw_source.name, (root / "wiki" / "log.md").read_text(encoding="utf-8"))

    def test_image_ingest_routes_to_assets_and_adds_preview(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            image = root / "diagram.jpg"
            image.write_bytes(b"jpeg-bytes")

            result = ingest_source(root, image, title="System Diagram")

            self.assertEqual(result["raw_source"].parent.name, "assets")
            page_text = result["wiki_page"].read_text(encoding="utf-8")
            self.assertIn('title: "System Diagram"', page_text)
            self.assertIn("![System Diagram]", page_text)
            self.assertIn("Kind: `asset`", page_text)

    def test_docx_ingest_extracts_markdown_into_source_page(self) -> None:
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not installed")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "brief.docx"
            doc = Document()
            doc.add_heading("Weekly Brief", level=1)
            doc.add_paragraph("This document explains the weekly operating rhythm.")
            doc.save(str(source))

            result = ingest_source(root, source, title="Weekly Brief")

            page_text = result["wiki_page"].read_text(encoding="utf-8")
            self.assertIn("## Extracted Content", page_text)
            self.assertIn("This document explains the weekly operating rhythm.", page_text)
            self.assertIn("# Weekly Brief", page_text)

    def test_xlsx_ingest_extracts_sheet_tables_into_source_page(self) -> None:
        try:
            from openpyxl import Workbook
        except ImportError:
            self.skipTest("openpyxl not installed")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "metrics.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Metrics"
            sheet.append(["Metric", "Value"])
            sheet.append(["Leads", 42])
            workbook.save(str(source))

            result = ingest_source(root, source, title="Metrics")

            page_text = result["wiki_page"].read_text(encoding="utf-8")
            self.assertIn("## Extracted Content", page_text)
            self.assertIn("## Sheet: Metrics", page_text)
            self.assertIn("| Metric | Value |", page_text)
            self.assertIn("| Leads | 42 |", page_text)

    def test_pdf_ingest_uses_extracted_text_for_source_page(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "paper.pdf"
            source.write_bytes(b"%PDF-1.4\n")

            with patch(
                "llm_wiki.workspace.extract_pdf_markdown",
                return_value="## Page 1\n\nThis paper introduces the main result.",
            ) as mock_extract:
                result = ingest_source(root, source, title="Research Paper")

            page_text = result["wiki_page"].read_text(encoding="utf-8")
            self.assertIn("## Extracted Content", page_text)
            self.assertIn("This paper introduces the main result.", page_text)
            mock_extract.assert_called_once()

    def test_ingest_url_creates_raw_snapshot_and_source_page(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            fake_source = WebSource(
                url="https://example.com/article",
                title="Example Article",
                markdown="This is the article body.",
                excerpt="This is the article body.",
                author="Test Author",
                published_date="2026-04-08",
                site_name="Example",
                fetched_at="2026-04-08 09:00 PDT",
            )

            with patch("llm_wiki.workspace.fetch_web_source", return_value=fake_source):
                result = ingest_url(root, "https://example.com/article")

            raw_source = result["raw_source"]
            wiki_page = result["wiki_page"]
            self.assertTrue(raw_source.exists())
            self.assertTrue(wiki_page.exists())
            raw_text = raw_source.read_text(encoding="utf-8")
            page_text = wiki_page.read_text(encoding="utf-8")
            self.assertIn('url: "https://example.com/article"', raw_text)
            self.assertIn("This is the article body.", raw_text)
            self.assertIn("[https://example.com/article](https://example.com/article)", page_text)
            self.assertIn("Author: `Test Author`", page_text)
            self.assertIn("Published: `2026-04-08`", page_text)

    def test_ingest_url_rejects_non_http_urls(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            with self.assertRaises(ValueError):
                ingest_url(root, "file:///tmp/secret.txt")

    def test_lint_reports_broken_links_orphans_and_missing_summaries(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)

            broken = root / "wiki" / "concepts" / "broken.md"
            broken.write_text(
                "# Broken\n\nHas a summary.\n\n## Related Pages\n\n- [Missing](missing.md)\n",
                encoding="utf-8",
            )

            orphan = root / "wiki" / "analyses" / "orphan.md"
            orphan.write_text("# Orphan\n\n- Only bullets here\n- Still no opening paragraph\n", encoding="utf-8")

            report = lint_workspace(root)

            self.assertEqual(len(report.broken_links), 1)
            self.assertTrue(any(page.path == broken.resolve() for page in report.orphan_pages))
            self.assertTrue(any(page.path == orphan.resolve() for page in report.missing_summaries))

    def test_heal_suggests_inbox_ingest_and_summary_repairs(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)

            (root / "raw" / "inbox" / "diagram.png").write_bytes(b"png")
            orphan = root / "wiki" / "analyses" / "orphan.md"
            orphan.write_text("# Orphan\n\n- no summary yet\n", encoding="utf-8")

            report = heal_workspace(root)
            messages = [item.message for item in report.suggestions]

            self.assertTrue(any("diagram.png" in message for message in messages))
            self.assertTrue(any("missing an opening summary paragraph" in message for message in messages))

    def test_compile_inbox_promotes_files_and_clears_inbox(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)

            note = root / "raw" / "inbox" / "research-note.txt"
            note.write_text("hello world", encoding="utf-8")
            diagram = root / "raw" / "inbox" / "diagram.png"
            diagram.write_bytes(b"png")

            report = compile_inbox(root)

            self.assertEqual(len(report.items), 2)
            self.assertFalse(note.exists())
            self.assertFalse(diagram.exists())
            self.assertTrue(any(item.raw_source.parent.name == "sources" for item in report.items))
            self.assertTrue(any(item.raw_source.parent.name == "assets" for item in report.items))
            self.assertTrue((root / "wiki" / "sources" / "research-note.md").exists())
            self.assertTrue((root / "wiki" / "sources" / "diagram.md").exists())

    def test_compile_inbox_respects_llmwikiignore(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)

            (root / ".llmwikiignore").write_text("raw/inbox/skip-me.txt\n", encoding="utf-8")
            keep = root / "raw" / "inbox" / "keep-me.txt"
            keep.write_text("keep", encoding="utf-8")
            skip = root / "raw" / "inbox" / "skip-me.txt"
            skip.write_text("skip", encoding="utf-8")

            report = compile_inbox(root)

            self.assertEqual(len(report.items), 1)
            self.assertEqual(report.items[0].original_path.name, "keep-me.txt")
            self.assertTrue(skip.exists())
            self.assertFalse((root / "wiki" / "sources" / "skip-me.md").exists())

    def test_heal_ignores_raw_files_listed_in_llmwikiignore(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)

            (root / ".llmwikiignore").write_text("raw/inbox/ignore.png\n", encoding="utf-8")
            (root / "raw" / "inbox" / "ignore.png").write_bytes(b"png")

            report = heal_workspace(root)

            self.assertFalse(any("ignore.png" in message.message for message in report.suggestions))

    def test_process_watch_changes_compiles_inbox_and_writes_status(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)
            inbox_file = root / "raw" / "inbox" / "note.txt"
            inbox_file.write_text("watch me", encoding="utf-8")

            report = process_watch_changes(
                root,
                [inbox_file],
                status_output_path=Path("outputs/status.md"),
            )

            self.assertEqual(len(report.compiled_items), 1)
            self.assertTrue(report.index_rebuilt)
            self.assertIsNotNone(report.status_output)
            assert report.status_output is not None
            self.assertTrue(report.status_output.exists())
            self.assertFalse(inbox_file.exists())
            self.assertTrue(any("Compiled 1 inbox file(s)." in message for message in report.messages))

    def test_process_watch_changes_rebuilds_index_for_wiki_edit(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)
            page = root / "wiki" / "concepts" / "systems.md"
            page.write_text("# Systems\n\nSystem-level thinking.\n", encoding="utf-8")

            report = process_watch_changes(root, [page])

            index_text = (root / "wiki" / "index.md").read_text(encoding="utf-8")
            self.assertTrue(report.index_rebuilt)
            self.assertIn("[Systems](concepts/systems.md)", index_text)
            self.assertTrue(any("Rebuilt wiki index after markdown changes." in message for message in report.messages))

    def test_process_watch_changes_ignores_generated_index_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)

            report = process_watch_changes(root, [root / "wiki" / "index.md"])

            self.assertFalse(report.has_work)
            self.assertEqual(report.messages, [])

    def test_build_brief_writes_markdown_report(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)

            concept = root / "wiki" / "concepts" / "attention.md"
            concept.write_text(
                "# Attention\n\n"
                "Mechanism for focusing on relevant context.\n\n"
                "## Key Points\n\n"
                "- Weights context dynamically\n"
                "- Used in transformers\n\n"
                "## Related Pages\n\n"
                "- [Overview](../overview.md)\n",
                encoding="utf-8",
            )

            analysis = root / "wiki" / "analyses" / "context-window.md"
            analysis.write_text(
                "# Context Window\n\n"
                "Maximum sequence length available to the model.\n\n"
                "## Key Points\n\n"
                "- Larger windows support more retrieval-free reasoning\n\n"
                "## Related Pages\n\n"
                "- [Attention](../concepts/attention.md)\n",
                encoding="utf-8",
            )

            report = build_brief(
                root,
                query="attention",
                page_refs=["analyses/context-window.md"],
                title="Attention Brief",
                output_path=Path("reports/attention-brief.md"),
            )

            self.assertIsNotNone(report.output_path)
            assert report.output_path is not None
            text = report.output_path.read_text(encoding="utf-8")
            self.assertIn("# Attention Brief", text)
            self.assertIn("Mechanism for focusing on relevant context.", text)
            self.assertIn("Used in transformers", text)
            self.assertIn("[Attention](../wiki/concepts/attention.md)", text)
            self.assertIn("[Context Window](../wiki/analyses/context-window.md)", text)

    def test_build_status_report_summarizes_workspace(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)
            create_page(root, "Attention", "concepts", "Mechanism for weighting relevant context.")
            (root / "raw" / "inbox" / "research-note.txt").write_text("attention note", encoding="utf-8")

            report = build_status_report(root, limit=3)

            self.assertIn("#", report.content)
            self.assertIn("## Corpus Check", report.content)
            self.assertIn("## Coverage", report.content)
            self.assertIn("## Suggested Actions", report.content)
            self.assertIn("raw/inbox/research-note.txt", report.content)
            self.assertIn("Citation gaps: 1", report.content)

    def test_build_status_report_writes_markdown_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)

            report = build_status_report(root, output_path=Path("outputs/status.md"))

            self.assertIsNotNone(report.output_path)
            assert report.output_path is not None
            self.assertTrue(report.output_path.exists())
            text = report.output_path.read_text(encoding="utf-8")
            self.assertIn("## Corpus Check", text)
            self.assertIn("## Structural Issues", text)

    def test_install_git_hooks_creates_managed_scripts_and_metadata(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)
            subprocess.run(["git", "init", "-q"], cwd=root, check=True)

            report = install_git_hooks(root, status_output_path=Path("outputs/status.md"))

            helper_text = report.helper_path.read_text(encoding="utf-8")
            pre_commit = next(path for path in report.hook_paths if path.name == "pre-commit")
            pre_commit_text = pre_commit.read_text(encoding="utf-8")
            metadata_text = report.metadata_path.read_text(encoding="utf-8")

            self.assertIn("llm-wiki-managed", helper_text)
            self.assertIn('git add -- "wiki/index.md"', helper_text)
            self.assertIn("LLM_WIKI_STAGE_STATUS='outputs/status.md'", helper_text)
            self.assertIn("outputs/status.md", metadata_text)
            self.assertIn("pythonpath", metadata_text)
            self.assertIn("llm-wiki-refresh", pre_commit_text)
            self.assertTrue(report.helper_path.stat().st_mode & 0o111)

    def test_install_git_hooks_requires_force_for_unmanaged_hook(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)
            subprocess.run(["git", "init", "-q"], cwd=root, check=True)

            hook_path = root / ".git" / "hooks" / "pre-commit"
            hook_path.parent.mkdir(parents=True, exist_ok=True)
            hook_path.write_text("#!/bin/sh\necho custom\n", encoding="utf-8")

            with self.assertRaises(RuntimeError):
                install_git_hooks(root)

            report = install_git_hooks(root, force=True)
            backups = [path for path in report.backup_paths if path.name.startswith("pre-commit.llmwiki.bak")]

            self.assertTrue(backups)
            self.assertIn("llm-wiki-managed", hook_path.read_text(encoding="utf-8"))

    def test_inspect_and_uninstall_git_hooks_reports_and_restores_backups(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)
            subprocess.run(["git", "init", "-q"], cwd=root, check=True)

            hook_path = root / ".git" / "hooks" / "pre-commit"
            hook_path.parent.mkdir(parents=True, exist_ok=True)
            hook_path.write_text("#!/bin/sh\necho custom\n", encoding="utf-8")
            install_git_hooks(root, status_output_path=Path("outputs/status.md"), force=True)

            status = inspect_git_hooks(root)
            uninstall = uninstall_git_hooks(root)

            self.assertTrue(status.is_installed)
            self.assertEqual(set(status.installed_hooks), {"pre-commit", "post-checkout", "post-merge"})
            self.assertEqual(status.status_output, "outputs/status.md")
            self.assertFalse((root / ".git" / "hooks" / "llm-wiki-refresh").exists())
            self.assertTrue(any(path.name == "pre-commit" for path in uninstall.restored_paths))
            self.assertIn("echo custom", hook_path.read_text(encoding="utf-8"))

    def test_git_hook_helper_supports_external_workspace_commit(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)
            subprocess.run(["git", "init", "-q"], cwd=root, check=True)
            subprocess.run(["git", "config", "user.email", "test@example.com"], cwd=root, check=True)
            subprocess.run(["git", "config", "user.name", "Test User"], cwd=root, check=True)

            page = root / "wiki" / "concepts" / "systems.md"
            page.write_text("# Systems\n\nSystem-level thinking.\n", encoding="utf-8")
            install_git_hooks(root, status_output_path=Path("outputs/status.md"))

            subprocess.run(["git", "add", "."], cwd=root, check=True)
            commit = subprocess.run(
                ["git", "commit", "-m", "smoke"],
                cwd=root,
                text=True,
                capture_output=True,
            )

            self.assertEqual(commit.returncode, 0, commit.stderr)
            self.assertTrue((root / "outputs" / "status.md").exists())
            index_text = (root / "wiki" / "index.md").read_text(encoding="utf-8")
            self.assertIn("[Systems](concepts/systems.md)", index_text)

    def test_git_hook_helper_skips_staging_ignored_status_output(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            init_workspace(root)
            subprocess.run(["git", "init", "-q"], cwd=root, check=True)
            subprocess.run(["git", "config", "user.email", "test@example.com"], cwd=root, check=True)
            subprocess.run(["git", "config", "user.name", "Test User"], cwd=root, check=True)
            (root / ".gitignore").write_text("outputs/*\n!outputs/.gitkeep\n", encoding="utf-8")

            page = root / "wiki" / "concepts" / "systems.md"
            page.write_text("# Systems\n\nSystem-level thinking.\n", encoding="utf-8")
            install_git_hooks(root, status_output_path=Path("outputs/status.md"))

            subprocess.run(["git", "add", "."], cwd=root, check=True)
            commit = subprocess.run(
                ["git", "commit", "-m", "smoke"],
                cwd=root,
                text=True,
                capture_output=True,
            )

            self.assertEqual(commit.returncode, 0, commit.stderr)
            self.assertTrue((root / "outputs" / "status.md").exists())
            tracked = subprocess.run(
                ["git", "ls-files", "outputs/status.md"],
                cwd=root,
                text=True,
                capture_output=True,
                check=True,
            )
            self.assertEqual(tracked.stdout.strip(), "")


if __name__ == "__main__":
    unittest.main()
